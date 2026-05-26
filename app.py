"""
Streamlit web interface for the AI Governance Navigator.
"""

from __future__ import annotations

import io
from functools import partial
from typing import BinaryIO

import anyio
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from pypdf import PdfReader

from pathlib import Path

from agent import GovernanceBrief, analyze_governance_question

PROJECT_ROOT = Path(__file__).resolve().parent

RISK_COLORS = {
    "High": ("#dc2626", "#fef2f2", "#991b1b"),
    "Medium": ("#ea580c", "#fff7ed", "#9a3412"),
    "Low": ("#16a34a", "#f0fdf4", "#166534"),
}


def _extract_pdf_text(uploaded_file: BinaryIO) -> str:
    reader = PdfReader(uploaded_file)
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)


def _build_question(question: str, artefacts: list[tuple[str, str]]) -> str:
    question = question.strip()
    if not artefacts:
        return question

    sections = [
        f"--- Uploaded artefact: {filename} ---\n{text.strip()}"
        for filename, text in artefacts
        if text.strip()
    ]
    if not sections:
        return question

    context = "\n\n".join(sections)
    return (
        "Additional context from uploaded internal artefacts:\n\n"
        f"{context}\n\n"
        "Governance question:\n"
        f"{question}"
    )


def _risk_badge_html(risk: str) -> str:
    border, background, text = RISK_COLORS.get(risk, RISK_COLORS["Medium"])
    return (
        f'<span style="background:{background};color:{text};border:1px solid {border};'
        f'padding:6px 14px;border-radius:999px;font-weight:600;">'
        f"Risk classification: {risk}</span>"
    )


def _render_brief(brief: GovernanceBrief) -> None:
    st.subheader("Question asked")
    st.write(brief.question)

    st.subheader("Jurisdictions consulted")
    if brief.jurisdictions_consulted:
        badges = " ".join(
            (
                '<span style="display:inline-block;background:#e8f0fe;color:#1a73e8;'
                "padding:4px 12px;border-radius:999px;margin:0 6px 6px 0;"
                f'font-size:0.9rem;">{jurisdiction}</span>'
            )
            for jurisdiction in brief.jurisdictions_consulted
        )
        st.markdown(badges, unsafe_allow_html=True)
    else:
        st.caption("No jurisdictions were consulted.")

    st.subheader("Key findings")
    if brief.key_findings:
        for jurisdiction, findings in brief.key_findings.items():
            with st.expander(jurisdiction, expanded=True):
                for finding in findings:
                    st.markdown(f"- {finding}")
    else:
        st.caption("No jurisdiction-specific findings were returned.")

    st.subheader("Convergences and divergences")
    col_convergences, col_divergences = st.columns(2)

    with col_convergences:
        st.markdown("**Convergences**")
        if brief.convergences:
            for item in brief.convergences:
                st.markdown(f"- {item}")
        else:
            st.caption("None identified.")

    with col_divergences:
        st.markdown("**Divergences**")
        if brief.divergences:
            for item in brief.divergences:
                st.markdown(f"- {item}")
        else:
            st.caption("None identified.")

    st.subheader("Risk classification")
    st.markdown(_risk_badge_html(brief.risk_classification), unsafe_allow_html=True)

    st.subheader("Sources")
    if brief.sources:
        for index, source in enumerate(brief.sources, start=1):
            st.markdown(f"{index}. {source}")
    else:
        st.caption("No sources were cited.")


def _build_word_document(brief: GovernanceBrief) -> bytes:
    doc = Document()

    title = doc.add_heading("AI Governance Brief", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_heading("Question asked", level=1)
    doc.add_paragraph(brief.question)

    doc.add_heading("Jurisdictions consulted", level=1)
    if brief.jurisdictions_consulted:
        doc.add_paragraph(", ".join(brief.jurisdictions_consulted))
    else:
        doc.add_paragraph("None")

    doc.add_heading("Key findings", level=1)
    for jurisdiction, findings in brief.key_findings.items():
        doc.add_heading(jurisdiction, level=2)
        for finding in findings:
            doc.add_paragraph(finding, style="List Bullet")

    doc.add_heading("Convergences", level=1)
    for item in brief.convergences:
        doc.add_paragraph(item, style="List Bullet")
    if not brief.convergences:
        doc.add_paragraph("None identified.")

    doc.add_heading("Divergences", level=1)
    for item in brief.divergences:
        doc.add_paragraph(item, style="List Bullet")
    if not brief.divergences:
        doc.add_paragraph("None identified.")

    doc.add_heading("Risk classification", level=1)
    risk_para = doc.add_paragraph()
    risk_run = risk_para.add_run(brief.risk_classification)
    risk_run.bold = True
    risk_run.font.size = Pt(14)
    if brief.risk_classification == "High":
        risk_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        risk_run.font.highlight_color = WD_COLOR_INDEX.RED
    elif brief.risk_classification == "Medium":
        risk_run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)
    else:
        risk_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    doc.add_heading("Sources", level=1)
    for index, source in enumerate(brief.sources, start=1):
        doc.add_paragraph(f"{index}. {source}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def main() -> None:
    load_dotenv()

    st.set_page_config(
        page_title="AI Governance Navigator",
        page_icon="⚖️",
        layout="wide",
    )

    if not (PROJECT_ROOT / "agent.py").exists() or not (PROJECT_ROOT / "server.py").exists():
        st.error(
            "This Streamlit session is not running the AI Governance Navigator project. "
            f"Expected `agent.py` and `server.py` in `{PROJECT_ROOT}`."
        )
        st.markdown(
            "Run from the project folder:\n\n"
            "```powershell\n"
            "cd \"C:\\Users\\nakul\\AI Stuff\\ai-governance-navigator\"\n"
            ".\\run_app.ps1\n"
            "```\n\n"
            "Or:\n\n"
            "```powershell\n"
            "streamlit run \"C:\\Users\\nakul\\AI Stuff\\ai-governance-navigator\\app.py\" --server.port 8502\n"
            "```"
        )
        st.stop()

    st.title("AI Governance Navigator")
    st.markdown(
        "Compare your AI governance questions against major regulatory frameworks:\n\n"
        "- **EU AI Act** (European Union)\n"
        "- **NIST AI RMF** (United States)\n"
        "- **MAS** AI guidelines (Singapore)\n"
        "- **UK AI Policy**\n"
        "- **FATF** guidance on AI and financial crime\n"
        "- **India** — NITI Aayog Responsible AI principles and the Digital Personal Data Protection Act\n\n"
        "Upload optional internal artefacts for additional context, then receive "
        "a structured brief with findings, convergences, divergences, and risk classification."
    )

    if not st.session_state.get("env_checked"):
        import os

        if not os.getenv("ANTHROPIC_API_KEY"):
            st.error("ANTHROPIC_API_KEY is not set. Add it to your `.env` file before running analyses.")
        st.session_state["env_checked"] = True

    with st.sidebar:
        st.header("Optional artefacts")
        st.caption(
            "Upload internal documents such as an AI policy, model card, or risk assessment. "
            "Their text will be included as additional context for the analysis."
        )
        uploaded_files = st.file_uploader(
            "Upload PDF artefacts",
            type=["pdf"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )

        if uploaded_files:
            st.markdown(f"**{len(uploaded_files)} file(s) attached**")
            for uploaded in uploaded_files:
                st.markdown(f"- {uploaded.name}")

    question = st.text_area(
        "Governance question",
        placeholder=(
            "Example: What are the requirements for deploying a high-risk AI credit scoring "
            "system across the EU and Singapore?"
        ),
        height=140,
    )

    submitted = st.button("Analyze governance question", type="primary", use_container_width=True)

    if submitted:
        if not question.strip():
            st.warning("Enter a governance question before submitting.")
            return

        artefacts: list[tuple[str, str]] = []
        for uploaded in uploaded_files or []:
            try:
                text = _extract_pdf_text(io.BytesIO(uploaded.getvalue()))
                if text.strip():
                    artefacts.append((uploaded.name, text))
                else:
                    st.warning(f"Could not extract text from `{uploaded.name}`.")
            except Exception as exc:
                st.warning(f"Failed to read `{uploaded.name}`: {exc}")

        full_question = _build_question(question, artefacts)

        try:
            with st.spinner("Analyzing governance question across regulatory frameworks..."):
                brief = anyio.run(partial(analyze_governance_question), full_question)
            st.session_state["brief"] = brief
            st.session_state["display_question"] = question.strip()
        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            return

    if "brief" in st.session_state:
        brief: GovernanceBrief = st.session_state["brief"]
        display_question = st.session_state.get("display_question", brief.question)
        brief_for_display = GovernanceBrief(
            question=display_question,
            jurisdictions_consulted=brief.jurisdictions_consulted,
            key_findings=brief.key_findings,
            convergences=brief.convergences,
            divergences=brief.divergences,
            risk_classification=brief.risk_classification,
            sources=brief.sources,
        )

        st.divider()
        _render_brief(brief_for_display)

        word_bytes = _build_word_document(brief_for_display)
        st.download_button(
            label="Export to Word",
            data=word_bytes,
            file_name="governance_brief.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=False,
        )


if __name__ == "__main__":
    main()
